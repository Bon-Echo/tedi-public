"""TDD DOCX generator with Bon Echo branding."""

import io
from typing import Any

import structlog
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

logger = structlog.get_logger(__name__)

# Bon Echo brand colours
_BONECHO_BLUE = RGBColor(0x1A, 0x56, 0xDB)   # #1A56DB
_FOOTER_TEXT = "© BonEcho · Confidential · tedi.bonecho.ai"


class TDDGenerator:
    """Generates a branded DOCX from a structured TDD dictionary."""

    def generate_docx(self, tdd: dict[str, Any]) -> bytes:
        """Build a DOCX document from the TDD dict.

        Args:
            tdd: Parsed TDD dict with keys: project_name, company_name,
                 project_overview, current_state, pain_points,
                 recommended_agents, integration_points, open_questions.

        Returns:
            Raw DOCX bytes.
        """
        doc = Document()

        self._configure_styles(doc)
        self._add_footer(doc)
        self._add_title_page(doc, tdd)
        doc.add_page_break()
        self._add_body(doc, tdd)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info("tdd_docx_generated", company=tdd.get("company_name", "unknown"))
        return buffer.read()

    @staticmethod
    def get_filename(tdd: dict[str, Any]) -> str:
        """Return the output filename for the DOCX.

        Args:
            tdd: TDD dict containing company_name.

        Returns:
            Filename string like ``Acme_Corp_TDD.docx``.
        """
        company = tdd.get("company_name") or "Unknown_Company"
        safe_name = company.replace(" ", "_").replace("/", "-")
        return f"{safe_name}_TDD.docx"

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _configure_styles(doc: Document) -> None:
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_footer(self, doc: Document) -> None:
        """Add the Bon Echo branding footer to the default section."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run(_FOOTER_TEXT)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)  # muted grey

        # Horizontal rule above footer via paragraph border
        self._add_top_border(para)

    @staticmethod
    def _add_top_border(para: Any) -> None:
        """Add a top border to a paragraph via OOXML."""
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "D1D5DB")
        pBdr.append(top)
        pPr.append(pBdr)

    def _add_title_page(self, doc: Document, tdd: dict[str, Any]) -> None:
        project_name = tdd.get("project_name") or "Untitled Project"
        company_name = tdd.get("company_name") or "Unknown Company"

        # Spacer
        doc.add_paragraph()
        doc.add_paragraph()

        # "TECHNICAL DESIGN DOCUMENT" — large title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("TECHNICAL DESIGN DOCUMENT")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = _BONECHO_BLUE
        run.font.name = "Calibri"

        # Company name
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = company_para.add_run(company_name)
        run.font.size = Pt(20)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

        # Project name
        project_para = doc.add_paragraph()
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = project_para.add_run(project_name)
        run.font.size = Pt(14)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        # Spacer
        doc.add_paragraph()

        # Prepared-by line
        prepared_para = doc.add_paragraph()
        prepared_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = prepared_para.add_run("Prepared by BonEcho · Tedi Discovery Agent")
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    def _add_body(self, doc: Document, tdd: dict[str, Any]) -> None:
        # 1. Project Overview
        self._add_section(doc, "1. Project Overview", tdd.get("project_overview", ""))

        # 2. Current State
        self._add_section(doc, "2. Current State", tdd.get("current_state", ""))

        # 3. Pain Points
        doc.add_heading("3. Pain Points", level=1)
        pain_points = tdd.get("pain_points") or []
        if pain_points:
            for point in pain_points:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(str(point))
        else:
            doc.add_paragraph("No pain points captured.")

        # 4. Recommended Agents
        doc.add_heading("4. Recommended Agents", level=1)
        agents = tdd.get("recommended_agents") or []
        if agents:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for cell, label in zip(hdr, ["Agent Name", "Purpose", "Priority"]):
                cell.text = label
                cell.paragraphs[0].runs[0].bold = True
            for agent in agents:
                row = table.add_row().cells
                row[0].text = str(agent.get("name", ""))
                row[1].text = str(agent.get("purpose", ""))
                row[2].text = str(agent.get("priority", ""))
        else:
            doc.add_paragraph("No recommended agents identified.")

        doc.add_paragraph()  # spacing after table

        # 5. Integration Points
        doc.add_heading("5. Integration Points", level=1)
        integrations = tdd.get("integration_points") or []
        if integrations:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for cell, label in zip(hdr, ["System", "Type", "Description"]):
                cell.text = label
                cell.paragraphs[0].runs[0].bold = True
            for intg in integrations:
                row = table.add_row().cells
                row[0].text = str(intg.get("system", ""))
                row[1].text = str(intg.get("type", ""))
                row[2].text = str(intg.get("description", ""))
        else:
            doc.add_paragraph("No integration points identified.")

        doc.add_paragraph()  # spacing after table

        # 6. Open Questions
        doc.add_heading("6. Open Questions", level=1)
        questions = tdd.get("open_questions") or []
        if questions:
            for i, question in enumerate(questions, 1):
                doc.add_paragraph(f"{i}. {question}")
        else:
            doc.add_paragraph("No open questions.")

    @staticmethod
    def _add_section(doc: Document, heading: str, content: str) -> None:
        doc.add_heading(heading, level=1)
        doc.add_paragraph(content if content else "No information captured for this section.")
